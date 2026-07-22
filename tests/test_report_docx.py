from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "simulator" / "versions" / "v2_postgres_react_ops" / "backend"
sys.path.insert(0, str(BACKEND))


def test_anomaly_report_docx_fills_template_without_visible_placeholders() -> None:
    from report_docx import render_anomaly_report_docx

    report = {
        "report_metadata": {
            "report_id": "REPORT-30",
            "generated_at": "2026-07-22T09:00:00+09:00",
            "source_card_id": "CARD-30",
        },
        "target_asset": {
            "asset_label": "30번 열교환 계통",
            "configuration_type": "열교환기",
            "window_start": "2023-03-12T14:00:00+09:00",
            "window_end": "2023-03-12T15:00:00+09:00",
        },
        "priority_summary": {
            "priority_level": "High",
            "priority_score": 87.5,
            "operator_review": "Required",
            "confidence": "medium",
            "urgency": "today",
            "priority_reason": "환수온도 저하와 열교환 효율 이상이 함께 확인되었습니다.",
        },
        "situation_summary": {
            "headline": "열교환기 운전 이상 검토 필요",
            "summary": "환수온도 저하와 열교환 효율 이상이 확인되어 현장 검토가 필요합니다.",
            "current_status": "운영자 검토 대기",
            "impact_summary": "난방 공급 안정성 저하 가능성이 있습니다.",
        },
        "key_evidence": [
            {
                "label": "두 모델 판단 불일치",
                "value": "Risk 높음 / 보조 판단 낮음",
                "interpretation": "계측값과 현장 상태를 함께 확인해야 합니다.",
                "confidence": "high",
                "evidence_ref_ids": ["ref-1"],
            }
        ],
        "risk_analysis": {
            "risk_level": "High",
            "risk_summary": "열교환 성능 저하가 지속될 가능성을 확인해야 합니다.",
        },
        "recommended_actions": [
            {
                "action": "환수온도와 열교환기 입출구 온도를 확인합니다.",
                "urgency": "today",
                "owner_hint": "현장 운영팀",
            }
        ],
        "evidence_references": [
            {
                "ref_id": "ref-1",
                "source_type": "ops_evidence",
                "title": "열교환기 운전 데이터",
                "excerpt": "환수온도 저하 확인",
            }
        ],
        "operator_note": {
            "note": "현장 확인 전 원인을 단정하지 않습니다.",
            "review_reasons": ["모델 판단 불일치"],
        },
    }

    rendered = render_anomaly_report_docx(
        report,
        alert_id="ALERT-30",
        building_name="범지기마을9단지한신휴플러스리버파크아파트",
        machine_room="기계실 30",
        status_label="검토 대기",
    )
    document = Document(BytesIO(rendered))
    full_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    assert "{{" not in full_text
    assert "작성 규칙:" not in full_text
    assert "REPORT-30" in full_text
    assert "범지기마을9단지한신휴플러스리버파크아파트" in full_text
    assert "열교환기" in full_text
    assert "환수온도 저하와 열교환 효율 이상" in full_text
    assert "ALERT-30" in full_text
    assert "원자료 없음" in full_text
    assert "5. 판정과 후속 조치" in full_text
    assert "불확실성 및 제외 조건" not in full_text
    assert "검토 및 승인" not in full_text
    assert "부록" not in full_text


def test_anomaly_report_docx_uses_selected_document_version() -> None:
    from report_docx import render_anomaly_report_docx

    rendered = render_anomaly_report_docx(
        {"report_metadata": {"report_id": "REPORT-31"}},
        alert_id=None,
        building_name="테스트 건물",
        machine_room="기계실 31",
        status_label="검토 중",
        document_version=3,
    )

    document = Document(BytesIO(rendered))
    full_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert "v3" in full_text
