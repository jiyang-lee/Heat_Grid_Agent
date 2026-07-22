from __future__ import annotations

from collections.abc import Mapping, Sequence
from io import BytesIO
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell


_TEMPLATE_PATH = Path(__file__).with_name("assets") / "report_template.docx"
_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


def render_anomaly_report_docx(
    report: Mapping[str, Any],
    *,
    alert_id: str | None,
    building_name: str,
    machine_room: str,
    status_label: str,
    document_version: int = 1,
) -> bytes:
    document = Document(_TEMPLATE_PATH)
    metadata = _mapping(report.get("report_metadata"))
    asset = _mapping(report.get("target_asset"))
    priority = _mapping(report.get("priority_summary"))
    situation = _mapping(report.get("situation_summary"))
    risk = _mapping(report.get("risk_analysis"))
    evidence = _mappings(report.get("key_evidence"))
    measurements = _mappings(report.get("sensor_measurements"))
    model_judgment = _mapping(report.get("model_judgment"))
    actions = _mappings(report.get("recommended_actions"))

    _fill_metadata(
        document.tables[0], metadata, asset, building_name, machine_room, status_label, document_version
    )
    _fill_decision_summary(document.tables[2], priority, situation, evidence)
    _set_cell(document.tables[3].cell(1, 0), _report_text(situation.get("summary")))
    _set_cell(document.tables[4].cell(1, 0), _actions_summary(actions))
    _fill_alert_overview(
        document.tables[5], metadata, asset, situation, alert_id, building_name, machine_room
    )
    _set_cell(
        document.tables[6].cell(0, 0),
        "운전 데이터 시계열 원자료가 제공된 경우 공급·환수 온도, 유량, 차압 추세를 표시합니다.",
    )
    _fill_sensor_table(document.tables[7], measurements, evidence)
    _fill_model_table(document.tables[8], priority, risk, evidence, model_judgment)
    _set_cell(document.tables[9].cell(1, 0), _disagreement_text(priority, evidence))
    _fill_decision_table(document.tables[11], actions)
    _remove_report_sections(document)
    _remove_remaining_placeholders(document)
    document.core_properties.title = "AI 이상 분석 보고서"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _fill_metadata(table, metadata, asset, building_name, machine_room, status_label, document_version) -> None:
    rows = (
        ("보고서 번호", metadata.get("report_id"), "보고서 유형", "AI 이상 분석"),
        ("대상 건물", building_name, "기계실", machine_room),
        ("대상 설비/계통", asset.get("asset_label") or asset.get("configuration_type"), "설비 ID", ""),
        (
            "대상 기간",
            f"{_text(asset.get('window_start'))} ~ {_text(asset.get('window_end'))}",
            "생성 일시",
            metadata.get("generated_at"),
        ),
        ("작성자", "AI 초안", "검토자", ""),
        ("승인 상태", status_label, "문서 버전", f"v{document_version}"),
    )
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            _set_cell(table.cell(row_index, column_index), value)


def _fill_decision_summary(table, priority, situation, evidence) -> None:
    disagreement = _find_evidence(evidence, "불일치")
    values = (
        situation.get("current_status") or situation.get("headline"),
        priority.get("priority_level"),
        priority.get("priority_score"),
        priority.get("confidence"),
        priority.get("urgency"),
        priority.get("operator_review"),
        "불일치" if disagreement else "일치 여부 확인 필요",
        "원자료 확인 필요",
    )
    for row_index in range(2):
        for column_index in range(4):
            _set_cell(table.cell(row_index * 2 + 1, column_index), values[row_index * 4 + column_index])


def _fill_alert_overview(table, metadata, asset, situation, alert_id, building_name, machine_room) -> None:
    values = (
        alert_id,
        metadata.get("source_card_id"),
        asset.get("window_start"),
        f"{_text(asset.get('window_start'))} ~ {_text(asset.get('window_end'))}",
        f"{building_name} / {machine_room}",
        asset.get("configuration_type") or asset.get("asset_label"),
        situation.get("current_status"),
        situation.get("impact_summary"),
    )
    value_index = 0
    for row_index in range(1, 5):
        for column_index in (1, 3):
            _set_cell(table.cell(row_index, column_index), values[value_index])
            value_index += 1


def _fill_sensor_table(table, measurements, evidence) -> None:
    labels = ("공급 온도", "환수 온도", "온도차 ΔT", "유량", "차압")
    for row_index, label in enumerate(labels, start=1):
        item = _find_measurement(measurements, label) or _find_evidence(evidence, label.replace(" ", ""))
        _set_cell(table.cell(row_index, 0), label)
        _set_cell(table.cell(row_index, 1), item.get("current_value", item.get("value", "")) if item else "")
        _set_cell(table.cell(row_index, 2), item.get("reference", "") if item else "")
        _set_cell(table.cell(row_index, 3), item.get("delta", "") if item else "")
        _set_cell(table.cell(row_index, 4), item.get("data_status", "확인됨") if item else "원자료 없음")
        _set_cell(table.cell(row_index, 5), item.get("judgement", item.get("interpretation", "확인 필요")) if item else "확인 필요")


def _fill_model_table(table, priority, risk, evidence, model_judgment) -> None:
    disagreement = _find_evidence(evidence, "불일치")
    rows = (
        ("Anomaly", model_judgment.get("anomaly_score", priority.get("priority_score")), model_judgment.get("anomaly_label", situation_label(risk)), model_judgment.get("reason", risk.get("risk_summary")), "모델 산출값"),
        ("Risk", risk.get("risk_score", priority.get("priority_score")), risk.get("risk_level"), risk.get("risk_summary"), "운영 영향 확인"),
        ("Lead time", priority.get("urgency"), priority.get("leadtime_bucket", "검토 필요"), priority.get("priority_reason"), "대응 시점 검토"),
        ("M1 Specialist", disagreement.get("value") if disagreement else model_judgment.get("m1_specialist_priority_score", "확인 필요"), "불일치" if disagreement else model_judgment.get("agreement", "확인 필요"), disagreement.get("interpretation") if disagreement else model_judgment.get("reason", "모델 비교 결과"), "비교 결과"),
        ("Hybrid Priority", priority.get("priority_score"), priority.get("priority_level"), priority.get("priority_reason"), "최종 검토"),
    )
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            _set_cell(table.cell(row_index, column_index), value)


def _fill_decision_table(table, actions) -> None:
    for row_index in range(1, 5):
        item = actions[row_index - 1] if row_index - 1 < len(actions) else {}
        _set_cell(table.cell(row_index, 2), "필요" if item else "미정")
        _set_cell(table.cell(row_index, 3), _report_text(item.get("action")))
        due_owner = " / ".join(filter(None, (_text(item.get("urgency")), _text(item.get("owner_hint")))))
        _set_cell(table.cell(row_index, 4), due_owner)


def _fill_approval_table(table, status_label) -> None:
    values = (
        ("작성", "AI 초안", "시스템", "분석 결과 기반 초안 생성", ""),
        ("검토", "", "운영 담당", "", ""),
        ("승인", "", "승인 담당", status_label, ""),
    )
    for row_index, row in enumerate(values, start=1):
        for column_index, value in enumerate(row):
            _set_cell(table.cell(row_index, column_index), value)


def _actions_summary(actions) -> str:
    if not actions:
        return "권고 조치가 없습니다."
    return "\n".join(f"{index}. {_text(item.get('action'))}" for index, item in enumerate(actions[:4], start=1))


def _disagreement_text(priority, evidence) -> str:
    item = _find_evidence(evidence, "불일치")
    return _text(item.get("interpretation")) or _text(priority.get("priority_reason"))


def _remove_report_sections(document: DocumentObject) -> None:
    # Keep the template's approved layout through section 5 and remove the unused tail.
    for table_index in (14, 13, 12, 10):
        table = document.tables[table_index]
        table._element.getparent().remove(table._element)
    for index in (26, 24, 22, 17):
        if index < len(document.paragraphs):
            paragraph = document.paragraphs[index]
            paragraph._element.getparent().remove(paragraph._element)
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("6."):
            paragraph.text = "5. 판정과 후속 조치"


def _find_measurement(items: Sequence[Mapping[str, Any]], label: str) -> Mapping[str, Any]:
    normalized = label.replace(" ", "").lower()
    for item in items:
        candidate = _text(item.get("label")).replace(" ", "").lower()
        if candidate == normalized:
            return item
    return {}


def situation_label(risk: Mapping[str, Any]) -> str:
    return _text(risk.get("risk_level")) or "확인 필요"


def _report_text(value: object) -> str:
    text = _text(value)
    text = re.sub(r"^\s*#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\s+(?=\d+[.)]\s)", "\n", text)
    text = re.sub(r"^\s*[-*]\s+", "• ", text, flags=re.MULTILINE)
    return text.strip()


def _find_evidence(items: Sequence[Mapping[str, Any]], keyword: str) -> Mapping[str, Any]:
    normalized = keyword.replace(" ", "")
    return next((item for item in items if normalized in _text(item.get("label")).replace(" ", "")), {})


def _source_type_label(value: object) -> str:
    return {
        "weather": "기상",
        "ops_evidence": "설비 정보",
        "rag_document": "전문 문서/RAG",
        "priority_card": "과거·운영 자료",
        "model_output": "모델 산출",
    }.get(_text(value), _text(value))


def _remove_remaining_placeholders(document: DocumentObject) -> None:
    for paragraph in document.paragraphs:
        _replace_paragraph_placeholders(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_paragraph_placeholders(paragraph)


def _replace_paragraph_placeholders(paragraph) -> None:
    for run in paragraph.runs:
        run.text = _PLACEHOLDER_RE.sub("", run.text)


def _set_cell(cell: _Cell, value: object) -> None:
    text = _report_text(value)
    paragraph = cell.paragraphs[0]
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    first_run.text = text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _mapping(value: object) -> Mapping[str, Any]:
    return value if isinstance(value, Mapping) else {}


def _mappings(value: object) -> tuple[Mapping[str, Any], ...]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes)):
        return ()
    return tuple(item for item in value if isinstance(item, Mapping))


def _strings(value: object) -> tuple[str, ...]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes)):
        return ()
    return tuple(_text(item) for item in value if _text(item))


def _text(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()
